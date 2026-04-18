"""Слоты тегов: apply-tag создаёт метаданные, revert восстанавливает исходный текст."""

import io
import re
import uuid
import zipfile

from docx import Document
from fastapi.testclient import TestClient

from app.docx_ops import build_docx_from_plain_text
from app.main import app, templates, template_versions


def _joined_w_text(xml: str) -> str:
    parts = re.findall(r"<w:t[^>]*>(.*?)</w:t>", xml)
    return "".join(parts)


def test_apply_tag_creates_slot_and_list_tag_slots() -> None:
    templates.clear()
    template_versions.clear()

    client = TestClient(app)
    r = client.post("/api/templates/bootstrap-empty", json={"name": "T"})
    data = r.json()
    tid, vid = data["templateId"], data["versionId"]

    raw = build_docx_from_plain_text("Контрагент: ООО Ромашка")
    client.post(
        f"/api/templates/{tid}/versions/{vid}/upload-docx",
        files={
            "file": (
                "t.docx",
                raw,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )

    apply_r = client.post(
        f"/api/templates/{tid}/versions/{vid}/apply-tag",
        json={
            "findText": "ООО Ромашка",
            "tagId": "buyer_name",
            "replaceAll": False,
            "occurrenceIndex": 0,
        },
    )
    assert apply_r.status_code == 200
    body = apply_r.json()
    assert body.get("ok") is True
    assert "tagSlotId" in body

    list_r = client.get(f"/api/templates/{tid}/versions/{vid}/tag-slots")
    assert list_r.status_code == 200
    slots = list_r.json()
    assert len(slots) == 1
    assert slots[0]["originalPlainText"] == "ООО Ромашка"
    assert slots[0]["currentOccurrenceIndex"] == 0
    assert "{{buyer_name}}" in slots[0]["currentTemplate"] or "buyer_name" in slots[0]["currentTemplate"]


def test_revert_tag_restores_original_plain_text() -> None:
    templates.clear()
    template_versions.clear()

    client = TestClient(app)
    r = client.post("/api/templates/bootstrap-empty", json={"name": "T"})
    data = r.json()
    tid, vid = data["templateId"], data["versionId"]

    raw = build_docx_from_plain_text("X")
    client.post(
        f"/api/templates/{tid}/versions/{vid}/upload-docx",
        files={
            "file": (
                "t.docx",
                raw,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )

    apply_r = client.post(
        f"/api/templates/{tid}/versions/{vid}/apply-tag",
        json={
            "findText": "X",
            "replacementTemplate": "{{field_1}}{{field_2}}",
            "replaceAll": False,
            "occurrenceIndex": 0,
        },
    )
    assert apply_r.status_code == 200
    slot_id = apply_r.json()["tagSlotId"]

    rev_r = client.post(
        f"/api/templates/{tid}/versions/{vid}/revert-tag",
        json={
            "tagSlotId": slot_id,
            "findText": "{{field_1}}{{field_2}}",
            "occurrenceIndex": 0,
        },
    )
    assert rev_r.status_code == 200

    got = client.get(f"/api/templates/{tid}/versions/{vid}/docx-file")
    z = zipfile.ZipFile(io.BytesIO(got.content))
    xml = z.read("word/document.xml").decode("utf-8")
    assert "X" in _joined_w_text(xml)
    assert "{{field_1}}" not in _joined_w_text(xml)

    list_r = client.get(f"/api/templates/{tid}/versions/{vid}/tag-slots")
    assert list_r.json() == []


def test_apply_tag_with_tag_slot_id_updates_slot_and_revert_ignores_stale_find_text() -> None:
    templates.clear()
    template_versions.clear()

    client = TestClient(app)
    r = client.post("/api/templates/bootstrap-empty", json={"name": "T"})
    data = r.json()
    tid, vid = data["templateId"], data["versionId"]

    raw = build_docx_from_plain_text("REPLACE_ME")
    client.post(
        f"/api/templates/{tid}/versions/{vid}/upload-docx",
        files={
            "file": (
                "t.docx",
                raw,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )

    a1 = client.post(
        f"/api/templates/{tid}/versions/{vid}/apply-tag",
        json={
            "findText": "REPLACE_ME",
            "replacementTemplate": "{{a}}",
            "replaceAll": False,
            "occurrenceIndex": 0,
        },
    )
    slot_id = a1.json()["tagSlotId"]

    a2 = client.post(
        f"/api/templates/{tid}/versions/{vid}/apply-tag",
        json={
            "findText": "{{a}}",
            "replacementTemplate": "{{a}}{{b}}",
            "replaceAll": False,
            "occurrenceIndex": 0,
            "tagSlotId": slot_id,
        },
    )
    assert a2.status_code == 200

    slots_after_update = client.get(f"/api/templates/{tid}/versions/{vid}/tag-slots")
    assert slots_after_update.status_code == 200
    assert slots_after_update.json()[0]["currentOccurrenceIndex"] == 0

    rev = client.post(
        f"/api/templates/{tid}/versions/{vid}/revert-tag",
        json={
            "tagSlotId": slot_id,
            "findText": "{{a}}",
            "occurrenceIndex": 0,
        },
    )
    assert rev.status_code == 200


def test_revert_uses_slot_occurrence_index_when_client_index_is_stale() -> None:
    templates.clear()
    template_versions.clear()

    client = TestClient(app)
    r = client.post("/api/templates/bootstrap-empty", json={"name": "T"})
    data = r.json()
    tid, vid = data["templateId"], data["versionId"]

    raw = build_docx_from_plain_text("A A")
    client.post(
        f"/api/templates/{tid}/versions/{vid}/upload-docx",
        files={
            "file": (
                "t.docx",
                raw,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )

    first = client.post(
        f"/api/templates/{tid}/versions/{vid}/apply-tag",
        json={
            "findText": "A",
            "replacementTemplate": "{{x}}",
            "replaceAll": False,
            "occurrenceIndex": 0,
        },
    )
    assert first.status_code == 200
    slot_id = first.json()["tagSlotId"]

    # Передаём намеренно неправильный occurrenceIndex: сервер должен использовать индекс из слота.
    rev = client.post(
        f"/api/templates/{tid}/versions/{vid}/revert-tag",
        json={
            "tagSlotId": slot_id,
            "findText": "{{x}}",
            "occurrenceIndex": 1,
        },
    )
    assert rev.status_code == 200


def test_insert_edit_delete_flow_for_identical_tokens_keeps_correct_slot_binding() -> None:
    templates.clear()
    template_versions.clear()

    client = TestClient(app)
    r = client.post("/api/templates/bootstrap-empty", json={"name": "T"})
    data = r.json()
    tid, vid = data["templateId"], data["versionId"]

    raw = build_docx_from_plain_text("A A")
    client.post(
        f"/api/templates/{tid}/versions/{vid}/upload-docx",
        files={
            "file": (
                "t.docx",
                raw,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )

    second_apply = client.post(
        f"/api/templates/{tid}/versions/{vid}/apply-tag",
        json={
            "findText": "A",
            "replacementTemplate": "{{x}}",
            "replaceAll": False,
            "occurrenceIndex": 1,
        },
    )
    assert second_apply.status_code == 200
    second_slot_id = second_apply.json()["tagSlotId"]

    first_apply = client.post(
        f"/api/templates/{tid}/versions/{vid}/apply-tag",
        json={
            "findText": "A",
            "replacementTemplate": "{{x}}",
            "replaceAll": False,
            "occurrenceIndex": 0,
        },
    )
    assert first_apply.status_code == 200
    first_slot_id = first_apply.json()["tagSlotId"]

    slots = client.get(f"/api/templates/{tid}/versions/{vid}/tag-slots")
    assert slots.status_code == 200
    by_id = {s["id"]: s for s in slots.json()}
    assert {by_id[first_slot_id]["currentOccurrenceIndex"], by_id[second_slot_id]["currentOccurrenceIndex"]} == {0, 1}

    edit_first = client.post(
        f"/api/templates/{tid}/versions/{vid}/apply-tag",
        json={
            "tagSlotId": first_slot_id,
            "findText": "{{x}}",
            "replacementTemplate": "{{x}}{{y}}",
            "replaceAll": False,
            "occurrenceIndex": 0,
        },
    )
    assert edit_first.status_code == 200

    revert_second = client.post(
        f"/api/templates/{tid}/versions/{vid}/revert-tag",
        json={
            "tagSlotId": second_slot_id,
            "findText": "{{x}}",
            # Имитируем устаревшее значение индекса с клиента.
            "occurrenceIndex": 0,
        },
    )
    assert revert_second.status_code == 200

    got = client.get(f"/api/templates/{tid}/versions/{vid}/docx-file")
    z = zipfile.ZipFile(io.BytesIO(got.content))
    xml = z.read("word/document.xml").decode("utf-8")
    rendered = _joined_w_text(xml)
    assert "{{x}}{{y}}" in rendered
    assert "A" in rendered

    slots_after = client.get(f"/api/templates/{tid}/versions/{vid}/tag-slots")
    assert slots_after.status_code == 200
    left_ids = {s["id"] for s in slots_after.json()}
    assert second_slot_id not in left_ids
    assert first_slot_id in left_ids


def test_edit_composite_tag_removes_soft_line_break_in_preview_docx() -> None:
    templates.clear()
    template_versions.clear()

    client = TestClient(app)
    r = client.post("/api/templates/bootstrap-empty", json={"name": "T"})
    data = r.json()
    tid, vid = data["templateId"], data["versionId"]

    raw = build_docx_from_plain_text("X")
    client.post(
        f"/api/templates/{tid}/versions/{vid}/upload-docx",
        files={
            "file": (
                "t.docx",
                raw,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )

    inserted = client.post(
        f"/api/templates/{tid}/versions/{vid}/apply-tag",
        json={
            "findText": "X",
            "replacementTemplate": "{{field_1}}\n{{field_2}}",
            "replaceAll": False,
            "occurrenceIndex": 0,
        },
    )
    assert inserted.status_code == 200
    slot_id = inserted.json()["tagSlotId"]

    edited = client.post(
        f"/api/templates/{tid}/versions/{vid}/apply-tag",
        json={
            "tagSlotId": slot_id,
            "findText": "{{field_1}}\n{{field_2}}",
            "replacementTemplate": "{{field_1}} {{field_2}}",
            "replaceAll": False,
            "occurrenceIndex": 0,
        },
    )
    assert edited.status_code == 200

    got = client.get(f"/api/templates/{tid}/versions/{vid}/docx-file")
    assert got.status_code == 200
    z = zipfile.ZipFile(io.BytesIO(got.content))
    xml = z.read("word/document.xml").decode("utf-8")
    rendered = _joined_w_text(xml)
    assert "{{field_1}} {{field_2}}" in rendered
    # Для превью это значит: перенос между плейсхолдерами убран.
    assert "<w:br" not in xml

    slots_after = client.get(f"/api/templates/{tid}/versions/{vid}/tag-slots")
    assert slots_after.status_code == 200
    assert slots_after.json()[0]["currentTemplate"] == "{{field_1}} {{field_2}}"


def test_large_formatted_docx_apply_then_edit_tag_on_first_word() -> None:
    templates.clear()
    template_versions.clear()

    client = TestClient(app)
    r = client.post("/api/templates/bootstrap-empty", json={"name": "T"})
    data = r.json()
    tid, vid = data["templateId"], data["versionId"]

    doc = Document()
    first_word = "FIRSTWORD"
    for i in range(1, 51):
        p = doc.add_paragraph()
        lead = first_word if i == 1 else f"W{i:02d}"
        r1 = p.add_run(f"{lead} ")
        r1.bold = i % 2 == 0
        r1.italic = i % 3 == 0
        r2 = p.add_run(f"line-{i:02d} ")
        r2.underline = i % 4 == 0
        r3 = p.add_run("formatted-tail")
        if i % 5 == 0:
            r3.bold = True
            r3.italic = True

    buf = io.BytesIO()
    doc.save(buf)
    raw = buf.getvalue()

    up = client.post(
        f"/api/templates/{tid}/versions/{vid}/upload-docx",
        files={
            "file": (
                "large_formatted.docx",
                raw,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert up.status_code == 200

    apply_r = client.post(
        f"/api/templates/{tid}/versions/{vid}/apply-tag",
        json={
            "findText": first_word,
            "replacementTemplate": "{{field_1}}",
            "replaceAll": False,
            "occurrenceIndex": 0,
        },
    )
    assert apply_r.status_code == 200
    slot_id = apply_r.json()["tagSlotId"]

    edit_r = client.post(
        f"/api/templates/{tid}/versions/{vid}/apply-tag",
        json={
            "tagSlotId": slot_id,
            "findText": "{{field_1}}",
            "replacementTemplate": "{{field_1}} {{field_2}}",
            "replaceAll": False,
            "occurrenceIndex": 0,
        },
    )
    assert edit_r.status_code == 200

    got = client.get(f"/api/templates/{tid}/versions/{vid}/docx-file")
    assert got.status_code == 200
    z = zipfile.ZipFile(io.BytesIO(got.content))
    xml = z.read("word/document.xml").decode("utf-8")
    rendered = _joined_w_text(xml)
    assert "{{field_1}} {{field_2}}" in rendered
    assert first_word not in rendered
    assert "line-50" in rendered
    assert "<w:b" in xml
    assert "<w:i" in xml


def test_edit_composite_duplicate_without_reselect_works_with_stale_slot_index() -> None:
    """
    Регрессия: пользователь жмёт «Редактировать» -> меняет шаблон -> «Вставить тег»
    без повторного выделения в превью.
    """
    templates.clear()
    template_versions.clear()

    client = TestClient(app)
    r = client.post("/api/templates/bootstrap-empty", json={"name": "T"})
    data = r.json()
    tid, vid = data["templateId"], data["versionId"]

    raw = build_docx_from_plain_text("X")
    client.post(
        f"/api/templates/{tid}/versions/{vid}/upload-docx",
        files={
            "file": (
                "t.docx",
                raw,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )

    inserted = client.post(
        f"/api/templates/{tid}/versions/{vid}/apply-tag",
        json={
            "findText": "X",
            "replacementTemplate": "{{field_1}}\n{{field_1}}",
            "replaceAll": False,
            "occurrenceIndex": 0,
        },
    )
    assert inserted.status_code == 200
    slot_id = inserted.json()["tagSlotId"]

    # Имитируем рассинхрон индекса в слоте (до фикса это приводило к 404).
    version = template_versions[uuid.UUID(vid)]
    slot = next(s for s in version["tag_slots"] if str(s["id"]) == slot_id)
    slot["current_occurrence_index"] = 99

    # "Без повторного выделения": findText пустой, работаем только по tagSlotId.
    edited = client.post(
        f"/api/templates/{tid}/versions/{vid}/apply-tag",
        json={
            "tagSlotId": slot_id,
            "findText": "",
            "replacementTemplate": "{{field_1}} {{field_1}}",
            "replaceAll": False,
            "occurrenceIndex": 0,
        },
    )
    assert edited.status_code == 200

    got = client.get(f"/api/templates/{tid}/versions/{vid}/docx-file")
    assert got.status_code == 200
    z = zipfile.ZipFile(io.BytesIO(got.content))
    xml = z.read("word/document.xml").decode("utf-8")
    rendered = _joined_w_text(xml)
    assert "{{field_1}} {{field_1}}" in rendered
    assert "<w:br" not in xml


def test_apply_tag_with_slot_id_allows_empty_find_text_for_edit_flow() -> None:
    templates.clear()
    template_versions.clear()

    client = TestClient(app)
    r = client.post("/api/templates/bootstrap-empty", json={"name": "T"})
    data = r.json()
    tid, vid = data["templateId"], data["versionId"]

    raw = build_docx_from_plain_text("X")
    client.post(
        f"/api/templates/{tid}/versions/{vid}/upload-docx",
        files={
            "file": (
                "t.docx",
                raw,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )

    inserted = client.post(
        f"/api/templates/{tid}/versions/{vid}/apply-tag",
        json={
            "findText": "X",
            "replacementTemplate": "{{field_1}}\\n {{field_1}}",
            "replaceAll": False,
            "occurrenceIndex": 0,
        },
    )
    assert inserted.status_code == 200
    slot_id = inserted.json()["tagSlotId"]

    edited = client.post(
        f"/api/templates/{tid}/versions/{vid}/apply-tag",
        json={
            "tagSlotId": slot_id,
            # В UI-режиме «Редактировать» findText может не отправляться.
            "findText": "",
            "replacementTemplate": "{{field_1}} {{field_1}}",
            "replaceAll": False,
            "occurrenceIndex": 0,
        },
    )
    assert edited.status_code == 200

    got = client.get(f"/api/templates/{tid}/versions/{vid}/docx-file")
    assert got.status_code == 200
    z = zipfile.ZipFile(io.BytesIO(got.content))
    xml = z.read("word/document.xml").decode("utf-8")
    rendered = _joined_w_text(xml)
    assert "{{field_1}} {{field_1}}" in rendered
    assert "<w:br" not in xml


def test_revert_resyncs_remaining_slot_indices_for_same_template() -> None:
    templates.clear()
    template_versions.clear()

    client = TestClient(app)
    r = client.post("/api/templates/bootstrap-empty", json={"name": "T"})
    data = r.json()
    tid, vid = data["templateId"], data["versionId"]

    raw = build_docx_from_plain_text("A A A")
    client.post(
        f"/api/templates/{tid}/versions/{vid}/upload-docx",
        files={
            "file": (
                "t.docx",
                raw,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )

    # 1) Метим первое A
    s1 = client.post(
        f"/api/templates/{tid}/versions/{vid}/apply-tag",
        json={
            "findText": "A",
            "replacementTemplate": "{{x}}",
            "replaceAll": False,
            "occurrenceIndex": 0,
        },
    )
    assert s1.status_code == 200
    slot1 = s1.json()["tagSlotId"]

    # 2) Метим третье исходное A (в текущем документе это occurrenceIndex=1 для findText=A)
    s2 = client.post(
        f"/api/templates/{tid}/versions/{vid}/apply-tag",
        json={
            "findText": "A",
            "replacementTemplate": "{{x}}",
            "replaceAll": False,
            "occurrenceIndex": 1,
        },
    )
    assert s2.status_code == 200
    slot2 = s2.json()["tagSlotId"]

    slots_before = client.get(f"/api/templates/{tid}/versions/{vid}/tag-slots")
    assert slots_before.status_code == 200
    by_id_before = {s["id"]: s for s in slots_before.json()}
    assert by_id_before[slot1]["currentOccurrenceIndex"] == 0
    assert by_id_before[slot2]["currentOccurrenceIndex"] == 1

    # Удаляем первый слот -> индекс второго должен автоматически пересчитаться (1 -> 0).
    rev = client.post(
        f"/api/templates/{tid}/versions/{vid}/revert-tag",
        json={
            "tagSlotId": slot1,
            "findText": "{{x}}",
            "occurrenceIndex": 0,
        },
    )
    assert rev.status_code == 200

    slots_after = client.get(f"/api/templates/{tid}/versions/{vid}/tag-slots")
    assert slots_after.status_code == 200
    assert len(slots_after.json()) == 1
    assert slots_after.json()[0]["id"] == slot2
    assert slots_after.json()[0]["currentOccurrenceIndex"] == 0
