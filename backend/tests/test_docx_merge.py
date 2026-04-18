import io
import re
import zipfile

from docx import Document

from app.docx_ops import (
    apply_docx_single_text_replacement,
    build_docx_from_plain_text,
    expand_replacement_escapes,
    merge_docx_placeholders,
)


def _joined_w_text(xml: str) -> str:
    parts = re.findall(r"<w:t[^>]*>(.*?)</w:t>", xml)
    return "".join(parts)


def test_merge_docx_replaces_placeholder_in_xml() -> None:
    body = "Hello {{seller_inn}} end."
    raw = build_docx_from_plain_text(body)
    name, out = merge_docx_placeholders(raw, {"seller_inn": "7701234567"})
    assert name.endswith(".docx")
    z = zipfile.ZipFile(io.BytesIO(out))
    xml = z.read("word/document.xml").decode("utf-8")
    assert "7701234567" in xml
    assert "{{seller_inn}}" not in xml


def test_merge_docx_replaces_split_run_placeholder() -> None:
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("Продавец ИНН: ")
    p.add_run("{{seller")
    p.add_run("_inn}}")
    buf = io.BytesIO()
    doc.save(buf)

    _, out = merge_docx_placeholders(buf.getvalue(), {"seller_inn": "7701234567"})
    z = zipfile.ZipFile(io.BytesIO(out))
    xml = z.read("word/document.xml").decode("utf-8")
    assert "7701234567" in _joined_w_text(xml)
    assert "{{seller" not in xml
    assert "_inn}}" not in xml


def test_merge_docx_keeps_table_and_run_formatting_tags() -> None:
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    p = table.cell(0, 0).paragraphs[0]
    p.add_run("ИНН: ")
    bold_run = p.add_run("{{seller")
    bold_run.bold = True
    p.add_run("_inn}}")
    buf = io.BytesIO()
    doc.save(buf)

    _, out = merge_docx_placeholders(buf.getvalue(), {"seller_inn": "7701234567"})
    z = zipfile.ZipFile(io.BytesIO(out))
    xml = z.read("word/document.xml").decode("utf-8")
    assert "<w:tbl" in xml
    assert "<w:b" in xml
    assert "7701234567" in _joined_w_text(xml)


def test_single_replacement_supports_line_breaks() -> None:
    raw = build_docx_from_plain_text("ООО Ромашка")
    out, ok = apply_docx_single_text_replacement(raw, "ООО Ромашка", "{{buyer_name}}\n{{buyer_inn}}", 0)
    assert ok is True
    z = zipfile.ZipFile(io.BytesIO(out))
    xml = z.read("word/document.xml").decode("utf-8")
    assert "<w:br" in xml
    assert "{{buyer_name}}" in xml
    assert "{{buyer_inn}}" in xml


def test_single_replacement_supports_paragraph_break_only_for_whole_paragraph() -> None:
    raw = build_docx_from_plain_text("ООО Ромашка")
    out, ok = apply_docx_single_text_replacement(raw, "ООО Ромашка", "{{buyer_name}}[[PARA_BREAK]]{{buyer_inn}}", 0)
    assert ok is True
    z = zipfile.ZipFile(io.BytesIO(out))
    xml = z.read("word/document.xml").decode("utf-8")
    assert xml.count("<w:p") >= 2

    raw2 = build_docx_from_plain_text("Контрагент: ООО Ромашка")
    _out2, ok2 = apply_docx_single_text_replacement(raw2, "ООО Ромашка", "{{buyer_name}}[[PARA_BREAK]]{{buyer_inn}}", 0)
    assert ok2 is False


def test_expand_replacement_escapes_turns_literal_backslash_sequences_into_chars() -> None:
    assert expand_replacement_escapes("a\\nb") == "a\nb"
    assert expand_replacement_escapes("a\\tb") == "a\tb"
    assert expand_replacement_escapes("a\\\\b") == "a\\b"
    assert expand_replacement_escapes("a\nb") == "a\nb"


def test_single_replacement_line_break_after_expand_literal_slash_n() -> None:
    raw = build_docx_from_plain_text("X")
    token = expand_replacement_escapes("111\\n222")
    out, ok = apply_docx_single_text_replacement(raw, "X", token, 0)
    assert ok is True
    z = zipfile.ZipFile(io.BytesIO(out))
    xml = z.read("word/document.xml").decode("utf-8")
    assert "<w:br" in xml
    assert "111" in _joined_w_text(xml)
    assert "222" in _joined_w_text(xml)


def test_cross_paragraph_composite_replace_removes_line_between_placeholders() -> None:
    """Два абзаца подряд: logical1 + \\n + logical2 == findText; замена без \\n склеивает в один абзац."""
    doc = Document()
    doc.add_paragraph("{{field_1}}")
    doc.add_paragraph("{{field_2}}")
    buf = io.BytesIO()
    doc.save(buf)
    raw = buf.getvalue()
    out, ok = apply_docx_single_text_replacement(
        raw, "{{field_1}}\n{{field_2}}", "{{field_1}}{{field_2}}", 0
    )
    assert ok is True
    z = zipfile.ZipFile(io.BytesIO(out))
    xml = z.read("word/document.xml").decode("utf-8")
    wt = _joined_w_text(xml)
    assert "{{field_1}}{{field_2}}" in wt
    assert "<w:br" not in xml


def test_single_replacement_find_text_with_newline_matches_w_br_paragraph() -> None:
    """findText как в предпросмотре (реальный \\n) должен находить абзац с w:br между плейсхолдерами."""
    raw = build_docx_from_plain_text("X")
    out, ok = apply_docx_single_text_replacement(raw, "X", "{{field_1}}\n{{field_2}}", 0)
    assert ok is True
    out2, ok2 = apply_docx_single_text_replacement(
        out, "{{field_1}}\n{{field_2}}", "{{field_1}} {{field_2}}", 0
    )
    assert ok2 is True
    z = zipfile.ZipFile(io.BytesIO(out2))
    xml = z.read("word/document.xml").decode("utf-8")
    assert "{{field_1}} {{field_2}}" in _joined_w_text(xml)
    assert "<w:br" not in xml


def test_merge_preserves_soft_break_between_two_placeholders() -> None:
    """После вставки шаблона с \\n между {{field_1}} и {{field_2}} merge не склеивает значения в одну строку."""
    raw = build_docx_from_plain_text("X")
    out, ok = apply_docx_single_text_replacement(raw, "X", "{{field_1}}\\n{{field_2}}", 0)
    assert ok is True
    _, merged = merge_docx_placeholders(out, {"field_1": "111", "field_2": "222"})
    z = zipfile.ZipFile(io.BytesIO(merged))
    xml = z.read("word/document.xml").decode("utf-8")
    assert "<w:br" in xml
    assert "{{field_1}}" not in xml
    assert "{{field_2}}" not in xml
    assert "111" in _joined_w_text(xml)
    assert "222" in _joined_w_text(xml)


