"""Работа с бинарным .docx: извлечение текста, сборка из текста, подстановка {{полей}} в OOXML."""

from __future__ import annotations

import io
import re
import zipfile
from datetime import datetime, timezone
from typing import Any
from xml.etree import ElementTree as ET
from xml.sax.saxutils import escape

from docx import Document

from app.generator import generate_docx


def xml_escape_for_word(text: str) -> str:
    return escape(text, entities={'"': "&quot;", "'": "&apos;"})


def extract_plain_text_from_docx(docx_bytes: bytes) -> str:
    doc = Document(io.BytesIO(docx_bytes))
    lines: list[str] = []
    for p in doc.paragraphs:
        lines.append(p.text)
    return "\n".join(lines)


def build_docx_from_plain_text(text: str) -> bytes:
    doc = Document()
    for line in text.split("\n"):
        doc.add_paragraph(line)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS = {"w": W_NS}
ET.register_namespace("w", W_NS)
PARA_BREAK_MARKER = "[[PARA_BREAK]]"


def norm_tag_fragment(s: str) -> str:
    """Согласованное сравнение фрагмента с предпросмотром (\\r vs \\n, пробелы после \\n перед {{)."""
    s = s.replace("\r\n", "\n").replace("\r", "\n")
    # NBSP и прочие Unicode-пробелы в превью, не только ASCII [ \\t]
    return re.sub(r"\n[^\S\r\n]+(?=\{\{)", "\n", s)


def _relaxed_tag_fragment_pattern(find_text: str) -> str:
    """Регулярное выражение: после каждого \\n допускаются пробелы/таб до следующего фрагмента."""
    ft = norm_tag_fragment(find_text)
    if "\n" not in ft:
        return re.escape(ft)
    parts = ft.split("\n")
    return r"\n\s*".join(re.escape(p) for p in parts)


def _find_all_spans(logical: str, find_text: str) -> list[tuple[int, int]]:
    """Все непересекающиеся вхождения find_text; при отсутствии точного совпадения — «мягкий» поиск между {{…}}."""
    if not find_text:
        return []
    spans: list[tuple[int, int]] = []
    pos = 0
    while True:
        i = logical.find(find_text, pos)
        if i < 0:
            break
        spans.append((i, i + len(find_text)))
        pos = i + len(find_text)
    if spans:
        return spans
    pat = _relaxed_tag_fragment_pattern(find_text)
    return [(m.start(), m.end()) for m in re.finditer(pat, logical)]


def expand_replacement_escapes(s: str) -> str:
    """
    Пользователь в поле «Шаблон вставки» часто вводит \\n и \\t как два символа.
    JSON с настоящим переводом строки приходит как символ \\n — его не трогаем.
    """
    out: list[str] = []
    i = 0
    while i < len(s):
        if s[i] == "\\" and i + 1 < len(s):
            nxt = s[i + 1]
            if nxt == "n":
                out.append("\n")
                i += 2
                continue
            if nxt == "t":
                out.append("\t")
                i += 2
                continue
            if nxt == "r":
                out.append("\r")
                i += 2
                continue
            if nxt == "\\":
                out.append("\\")
                i += 2
                continue
        out.append(s[i])
        i += 1
    return "".join(out)


def _word_xml_sort_key(fn: str) -> tuple[int, str]:
    """Стабильный порядок: тело документа раньше колонтитулов (как в предпросмотре)."""
    base = fn.rsplit("/", 1)[-1].lower()
    if base == "document.xml":
        return (0, base)
    if base.startswith("header"):
        return (1, base)
    if base.startswith("footer"):
        return (2, base)
    return (3, base)


def _total_in_para_occurrences_in_root(root: ET.Element, find_text: str) -> int:
    n = 0
    for p in root.findall(".//w:p", NS):
        logical = _paragraph_logical_text(p)
        n += len(_find_all_spans(logical, find_text))
    return n


def _cross_paragraph_pair_match_indices(root: ET.Element, find_text: str) -> list[int]:
    """
    Индексы i такие, что logical(p[i]) + '\\n' + logical(p[i+1]) совпадает с find_text
    (два подряд абзаца в XML-порядке = перенос между строками в предпросмотре).
    """
    paragraphs = root.findall(".//w:p", NS)
    matches: list[int] = []
    for i in range(len(paragraphs) - 1):
        l1 = _paragraph_logical_text(paragraphs[i])
        l2 = _paragraph_logical_text(paragraphs[i + 1])
        combined = l1 + "\n" + l2
        relaxed = l1.rstrip() + "\n" + l2.lstrip()
        if combined == find_text or relaxed == find_text:
            matches.append(i)
            continue
        spans = _find_all_spans(combined, find_text)
        if len(spans) == 1 and spans[0] == (0, len(combined)):
            matches.append(i)
    return matches


def _whole_paragraph_matches_find_text(logical: str, find_text: str) -> bool:
    spans = _find_all_spans(logical, find_text)
    return len(spans) == 1 and spans[0] == (0, len(logical))


def _replace_nth_substring(haystack: str, needle: str, repl: str, n: int) -> str | None:
    """Заменяет n-е вхождение needle (0-based) в haystack на repl. None, если вхождений меньше n + 1."""
    if not needle:
        return None
    spans = _find_all_spans(haystack, needle)
    if n >= len(spans):
        return None
    start, end = spans[n]
    return haystack[:start] + repl + haystack[end:]


def _w_tag(name: str) -> str:
    return f"{{{W_NS}}}{name}"


def _paragraph_text(p: ET.Element) -> str:
    texts = list(p.findall(".//w:t", NS))
    return "".join(t.text or "" for t in texts)


def _paragraph_logical_text(p: ET.Element) -> str:
    """Текст абзаца с учётом мягких переносов (w:br) как символов \\n — для подстановки полей без потери разрыва."""
    parts: list[str] = []
    for r in p.findall(".//w:r", NS):
        for sub in r:
            if sub.tag == _w_tag("t"):
                parts.append(sub.text or "")
            elif sub.tag == _w_tag("br"):
                parts.append("\n")
    return "".join(parts)


def _apply_flat_updated_to_wt_nodes(p: ET.Element, updated: str) -> bool:
    """Распределяет строку без \\n по существующим w:t (структура без разрывов)."""
    texts = list(p.findall(".//w:t", NS))
    if not texts:
        return False
    original = "".join(t.text or "" for t in texts)
    if updated == original:
        return False
    cursor = 0
    for idx, node in enumerate(texts):
        node_len = len(node.text or "")
        if idx == len(texts) - 1:
            node.text = updated[cursor:]
        else:
            node.text = updated[cursor : cursor + node_len]
        cursor += node_len
    return True


def _merge_all_placeholder_tokens_in_paragraph(p: ET.Element, values: dict[str, str]) -> bool:
    """
    Заменяет все {{id}} в абзаце за один проход по логической строке (с \\n для w:br),
    чтобы перенос между плейсхолдерами не терялся при подстановке значений.
    """
    logical = _paragraph_logical_text(p)
    if not logical or "{{" not in logical:
        return False
    updated = logical
    for src, val in values.items():
        if src in updated:
            updated = updated.replace(src, "" if val is None else str(val))
    if updated == logical:
        return False
    if "\n" in logical or "\n" in updated:
        _set_paragraph_text_with_breaks(p, updated)
        return True
    return _apply_flat_updated_to_wt_nodes(p, updated)


def _make_run_with_text(text: str, template_run: ET.Element | None, break_before: bool = False) -> ET.Element:
    run = ET.Element(_w_tag("r"))
    if template_run is not None:
        rpr = template_run.find("w:rPr", NS)
        if rpr is not None:
            run.append(ET.fromstring(ET.tostring(rpr, encoding="unicode")))
    if break_before:
        run.append(ET.Element(_w_tag("br")))
    text_el = ET.Element(_w_tag("t"))
    if text.startswith(" ") or text.endswith(" "):
        text_el.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    text_el.text = text
    run.append(text_el)
    return run


def _set_paragraph_text_with_breaks(p: ET.Element, text: str) -> None:
    template_run = p.find(".//w:r", NS)
    keep: list[ET.Element] = []
    for child in list(p):
        if child.tag == _w_tag("pPr"):
            keep.append(child)
        p.remove(child)
    for child in keep:
        p.append(child)
    lines = text.split("\n")
    if not lines:
        p.append(_make_run_with_text("", template_run))
        return
    p.append(_make_run_with_text(lines[0], template_run))
    for line in lines[1:]:
        p.append(_make_run_with_text(line, template_run, break_before=True))


def _replace_tokens_in_paragraph_runs(p: ET.Element, values: dict[str, str]) -> bool:
    texts = list(p.findall(".//w:t", NS))
    if not texts:
        return False
    original = "".join(t.text or "" for t in texts)
    updated = original
    for src, val in values.items():
        updated = updated.replace(src, val)
    if updated == original:
        return False

    cursor = 0
    for idx, node in enumerate(texts):
        node_len = len(node.text or "")
        if idx == len(texts) - 1:
            node.text = updated[cursor:]
        else:
            node.text = updated[cursor : cursor + node_len]
        cursor += node_len

    return True


def _replace_template_in_paragraph_runs(p: ET.Element, find_text: str, replace_template: str, replace_all: bool) -> tuple[bool, bool]:
    original = _paragraph_text(p)
    if not original:
        return False, False
    updated = original.replace(find_text, replace_template) if replace_all else original.replace(find_text, replace_template, 1)
    if updated == original:
        return False, False
    if "\n" in replace_template and PARA_BREAK_MARKER not in replace_template:
        _set_paragraph_text_with_breaks(p, updated)
        return True, False
    return _replace_tokens_in_paragraph_runs(p, {find_text: replace_template}), False


def _replace_exact_paragraph_with_parts(parent: ET.Element, p: ET.Element, replacement: str) -> bool:
    parts = replacement.split(PARA_BREAK_MARKER)
    if len(parts) <= 1:
        return False
    idx = list(parent).index(p)
    template_xml = ET.tostring(p, encoding="unicode")
    parent.remove(p)
    for offset, part in enumerate(parts):
        clone = ET.fromstring(template_xml)
        _set_paragraph_text_with_breaks(clone, part)
        parent.insert(idx + offset, clone)
    return True


def _replace_tokens_in_word_xml(xml_bytes: bytes, values: dict[str, str]) -> bytes:
    try:
        root = ET.fromstring(xml_bytes)
    except ET.ParseError:
        s = xml_bytes.decode("utf-8", errors="ignore")
        for key, val in values.items():
            token = "{{" + key + "}}"
            s = s.replace(token, xml_escape_for_word(val))
        return s.encode("utf-8")

    changed = False
    parent_map = {child: parent for parent in root.iter() for child in parent}
    for p in root.findall(".//w:p", NS):
        para_handled = False
        for src, val in values.items():
            if PARA_BREAK_MARKER in val:
                original = _paragraph_text(p)
                if original != src:
                    continue
                parent = parent_map.get(p)
                if parent is None:
                    continue
                changed = _replace_exact_paragraph_with_parts(parent, p, val) or changed
                para_handled = True
                break
        if para_handled:
            continue
        if _merge_all_placeholder_tokens_in_paragraph(p, values):
            changed = True
    if not changed:
        return xml_bytes
    return ET.tostring(root, encoding="utf-8", xml_declaration=True)


def apply_docx_text_replacements(docx_bytes: bytes, replacements: dict[str, str]) -> bytes:
    """Заменяет произвольные текстовые фрагменты в word/*.xml с сохранением OOXML-структуры."""
    replacements = {k: expand_replacement_escapes(v) for k, v in replacements.items()}
    buf_in = io.BytesIO(docx_bytes)
    buf_out = io.BytesIO()
    with (
        zipfile.ZipFile(buf_in, "r") as zin,
        zipfile.ZipFile(buf_out, "w", zipfile.ZIP_DEFLATED) as zout,
    ):
        for zinfo in zin.infolist():
            raw = zin.read(zinfo.filename)
            fn = zinfo.filename.replace("\\", "/")
            if fn.endswith(".xml") and fn.startswith("word/"):
                raw = _replace_tokens_in_word_xml(raw, replacements)
            zout.writestr(fn, raw, compress_type=zipfile.ZIP_DEFLATED)
    return buf_out.getvalue()


def _replace_nth_in_paragraph_runs(p: ET.Element, src: str, dst: str, target_index: int, seen: int) -> tuple[bool, int]:
    texts = list(p.findall(".//w:t", NS))
    if not texts:
        return False, seen
    original = "".join(t.text or "" for t in texts)
    pos = 0
    local_occurs: list[int] = []
    while True:
        i = original.find(src, pos)
        if i < 0:
            break
        local_occurs.append(i)
        pos = i + len(src)
    if not local_occurs:
        return False, seen

    if target_index < seen or target_index >= seen + len(local_occurs):
        return False, seen + len(local_occurs)

    idx_in_local = target_index - seen
    start = local_occurs[idx_in_local]
    updated = original[:start] + dst + original[start + len(src) :]

    cursor = 0
    for idx, node in enumerate(texts):
        node_len = len(node.text or "")
        if idx == len(texts) - 1:
            node.text = updated[cursor:]
        else:
            node.text = updated[cursor : cursor + node_len]
        cursor += node_len
    return True, seen + len(local_occurs)


def apply_docx_single_text_replacement(docx_bytes: bytes, find_text: str, replace_with: str, occurrence_index: int) -> tuple[bytes, bool]:
    """
    Заменяет только одно (N-е) вхождение текста в word/*.xml, сохраняя OOXML-структуру.
    occurrence_index — 0-based, **глобально по документу**:
    сначала все вхождения внутри одного w:p (файлы в порядке document → header → footer),
    затем межабзацные пары (тот же порядок файлов, пары сверху вниз).
    """
    replace_with = expand_replacement_escapes(replace_with)
    buf_in = io.BytesIO(docx_bytes)
    infos: list[zipfile.ZipInfo] = []
    raw_by_fn: dict[str, bytes] = {}
    with zipfile.ZipFile(buf_in, "r") as zin:
        infos = zin.infolist()
        for zi in infos:
            fn = zi.filename.replace("\\", "/")
            raw_by_fn[fn] = zin.read(zi.filename)

    word_xml_fns = sorted(
        [fn for fn in raw_by_fn if fn.startswith("word/") and fn.endswith(".xml")],
        key=_word_xml_sort_key,
    )
    parsed: dict[str, ET.Element | None] = dict.fromkeys(word_xml_fns)
    for fn in word_xml_fns:
        try:
            parsed[fn] = ET.fromstring(raw_by_fn[fn])
        except ET.ParseError:
            parsed[fn] = None

    total_in_para = 0
    for fn in word_xml_fns:
        root = parsed.get(fn)
        if root is not None:
            total_in_para += _total_in_para_occurrences_in_root(root, find_text)

    modified_fn: str | None = None
    replaced_any = False
    global_seen = 0

    for fn in word_xml_fns:
        root = parsed.get(fn)
        if root is None:
            continue
        parent_map = {child: parent for parent in root.iter() for child in parent}
        for p in root.findall(".//w:p", NS):
            logical = _paragraph_logical_text(p)
            spans_here = _find_all_spans(logical, find_text)
            local_count = len(spans_here)
            if local_count == 0:
                continue
            if occurrence_index < global_seen or occurrence_index >= global_seen + local_count:
                global_seen += local_count
                continue
            local_target = occurrence_index - global_seen
            if PARA_BREAK_MARKER in replace_with:
                if local_count != 1 or local_target != 0 or not _whole_paragraph_matches_find_text(logical, find_text):
                    break
                parent = parent_map.get(p)
                if parent is None:
                    break
                if _replace_exact_paragraph_with_parts(parent, p, replace_with):
                    modified_fn = fn
                    replaced_any = True
                break
            updated = _replace_nth_substring(logical, find_text, replace_with, local_target)
            if updated is None:
                global_seen += local_count
                continue
            if "\n" in logical or "\n" in updated:
                _set_paragraph_text_with_breaks(p, updated)
            else:
                _apply_flat_updated_to_wt_nodes(p, updated)
            modified_fn = fn
            replaced_any = True
            break
        if replaced_any:
            break

    if not replaced_any and "\n" in find_text:
        cross_idx = occurrence_index - total_in_para
        if cross_idx >= 0:
            cross_matches: list[tuple[str, int]] = []
            for fn in word_xml_fns:
                root = parsed.get(fn)
                if root is None:
                    continue
                for i0 in _cross_paragraph_pair_match_indices(root, find_text):
                    cross_matches.append((fn, i0))
            if cross_idx < len(cross_matches):
                fn, i0 = cross_matches[cross_idx]
                root = parsed[fn]
                if root is not None:
                    parent_map = {child: parent for parent in root.iter() for child in parent}
                    paragraphs = root.findall(".//w:p", NS)
                    p_a = paragraphs[i0]
                    p_b = paragraphs[i0 + 1]
                    if PARA_BREAK_MARKER not in replace_with:
                        _set_paragraph_text_with_breaks(p_a, replace_with)
                        parent = parent_map.get(p_b)
                        if parent is not None:
                            parent.remove(p_b)
                        modified_fn = fn
                        replaced_any = True

    buf_out = io.BytesIO()
    with zipfile.ZipFile(buf_out, "w", zipfile.ZIP_DEFLATED) as zout:
        for zi in infos:
            fn = zi.filename.replace("\\", "/")
            raw = raw_by_fn[fn]
            if replaced_any and modified_fn == fn and parsed.get(fn) is not None:
                raw = ET.tostring(parsed[fn], encoding="utf-8", xml_declaration=True)
            zout.writestr(zi, raw, compress_type=zipfile.ZIP_DEFLATED)

    return buf_out.getvalue(), replaced_any


def _normalize_fragment_for_matching(s: str) -> str:
    s = norm_tag_fragment(s)
    s = re.sub(r"[^\S\r\n]*\n[^\S\r\n]*", "\n", s)
    s = re.sub(r"\n{2,}", "\n", s)
    return s.strip()


def _remove_nth_large_fragment_in_word_xml(xml_bytes: bytes, find_text: str, occurrence_index: int) -> tuple[bytes, bool]:
    try:
        root = ET.fromstring(xml_bytes)
    except ET.ParseError:
        return xml_bytes, False

    paragraphs = root.findall(".//w:p", NS)
    if not paragraphs:
        return xml_bytes, False

    para_norm: list[str] = [_normalize_fragment_for_matching(_paragraph_logical_text(p)) for p in paragraphs]
    joined = "\n".join(para_norm)
    needle = _normalize_fragment_for_matching(find_text)
    if not joined or not needle:
        return xml_bytes, False

    spans = _find_all_spans(joined, needle)
    if occurrence_index < 0 or occurrence_index >= len(spans):
        return xml_bytes, False
    target_start, target_end = spans[occurrence_index]

    para_char_spans: list[tuple[int, int]] = []
    cursor = 0
    for idx, txt in enumerate(para_norm):
        start = cursor
        end = start + len(txt)
        para_char_spans.append((start, end))
        cursor = end + (1 if idx < len(para_norm) - 1 else 0)

    remove_indices: list[int] = []
    for i, (start, end) in enumerate(para_char_spans):
        intersects = start < target_end and end > target_start
        if intersects:
            remove_indices.append(i)
    if not remove_indices:
        return xml_bytes, False

    parent_map = {child: parent for parent in root.iter() for child in parent}
    removed_any = False
    for i in reversed(remove_indices):
        p = paragraphs[i]
        parent = parent_map.get(p)
        if parent is None:
            continue
        parent.remove(p)
        removed_any = True
    if not removed_any:
        return xml_bytes, False
    return ET.tostring(root, encoding="utf-8", xml_declaration=True), True


def remove_docx_fragment(docx_bytes: bytes, find_text: str, occurrence_index: int) -> tuple[bytes, bool]:
    """
    Удаляет N-е вхождение фрагмента:
    1) сначала точечной заменой (работает для небольших фрагментов),
    2) затем fallback для больших многоабзацных фрагментов.
    """
    updated, ok = apply_docx_single_text_replacement(docx_bytes, find_text, "", occurrence_index)
    if ok:
        return updated, True

    buf_in = io.BytesIO(docx_bytes)
    infos: list[zipfile.ZipInfo] = []
    raw_by_fn: dict[str, bytes] = {}
    with zipfile.ZipFile(buf_in, "r") as zin:
        infos = zin.infolist()
        for zi in infos:
            fn = zi.filename.replace("\\", "/")
            raw_by_fn[fn] = zin.read(zi.filename)

    word_xml_fns = sorted(
        [fn for fn in raw_by_fn if fn.startswith("word/") and fn.endswith(".xml")],
        key=_word_xml_sort_key,
    )

    global_seen = 0
    modified_fn: str | None = None
    for fn in word_xml_fns:
        raw = raw_by_fn[fn]
        try:
            root = ET.fromstring(raw)
        except ET.ParseError:
            continue
        paragraphs = root.findall(".//w:p", NS)
        if not paragraphs:
            continue
        para_norm = [_normalize_fragment_for_matching(_paragraph_logical_text(p)) for p in paragraphs]
        joined = "\n".join(para_norm)
        needle = _normalize_fragment_for_matching(find_text)
        if not joined or not needle:
            continue
        spans = _find_all_spans(joined, needle)
        local_count = len(spans)
        if local_count == 0:
            continue
        if occurrence_index < global_seen or occurrence_index >= global_seen + local_count:
            global_seen += local_count
            continue
        local_idx = occurrence_index - global_seen
        updated_xml, local_ok = _remove_nth_large_fragment_in_word_xml(raw, find_text, local_idx)
        if not local_ok:
            return docx_bytes, False
        raw_by_fn[fn] = updated_xml
        modified_fn = fn
        break

    if modified_fn is None:
        return docx_bytes, False

    buf_out = io.BytesIO()
    with zipfile.ZipFile(buf_out, "w", zipfile.ZIP_DEFLATED) as zout:
        for zi in infos:
            fn = zi.filename.replace("\\", "/")
            raw = raw_by_fn[fn]
            zout.writestr(zi, raw, compress_type=zipfile.ZIP_DEFLATED)
    return buf_out.getvalue(), True


def merge_docx_placeholders(docx_bytes: bytes, values: dict[str, str]) -> tuple[str, bytes]:
    """
    Подставляет значения в плейсхолдеры вида {{field_id}} во всех word/*.xml.
    Ограничение MVP: плейсхолдер должен целиком попадать в один XML-текстовый фрагмент Word
    (при разбиении Word на «раны» подстановка может не сработать — тогда вставьте поле заново одним фрагментом).
    """
    buf_in = io.BytesIO(docx_bytes)
    buf_out = io.BytesIO()
    token_replacements = {"{{" + key + "}}": val for key, val in values.items()}
    merged = apply_docx_text_replacements(docx_bytes, token_replacements)
    ts = datetime.now(timezone.utc).strftime("%Y%m%d%H%M%S")
    return f"dkp-{ts}.docx", merged


def apply_docx_conditional_blocks(docx_bytes: bytes, payload: dict[str, Any], blocks: list[dict[str, Any]]) -> tuple[bytes, list[str]]:
    """
    Применяет условные блоки к DOCX:
    - branch=if: блок видим, если payload[field] == equals_value (case-insensitive)
    - branch=else: блок видим, если условие НЕ выполняется
    Невидимые блоки удаляются заменой на пустую строку по (find_template, occurrence_index).
    """
    if not blocks:
        return docx_bytes, []
    out = docx_bytes
    warnings: list[str] = []

    by_template: dict[str, list[dict[str, Any]]] = {}
    for b in blocks:
        tpl = str(b.get("find_template") or "")
        if not tpl:
            continue
        by_template.setdefault(tpl, []).append(b)

    for template, group in by_template.items():
        to_delete = []
        for b in group:
            field = str(b.get("condition_field") or "")
            equals_value = str(b.get("equals_value") or "")
            branch = str(b.get("branch") or "if").lower()
            incoming = payload.get(field)
            cond_ok = incoming is not None and str(incoming).lower() == equals_value.lower()
            visible = cond_ok if branch != "else" else (not cond_ok)
            if visible:
                continue
            occ = b.get("occurrence_index")
            if not isinstance(occ, int) or occ < 0:
                warnings.append(f"block:{b.get('id')} invalid occurrence_index")
                continue
            to_delete.append((occ, b))

        # Удаляем с конца, чтобы не сдвигать индексы оставшихся вхождений того же шаблона.
        to_delete.sort(key=lambda x: x[0], reverse=True)
        for occ, b in to_delete:
            updated, ok = remove_docx_fragment(out, template, occ)
            if ok:
                out = updated
            else:
                warnings.append(f"block:{b.get('id')} stale or not found")
    return out, warnings


def render_version_to_docx(
    *,
    docx_bytes: bytes | None,
    docx_template_body: str,
    bindings_json: str,
    rules_json: str,
    payload_json: str,
    conditional_blocks: list[dict[str, Any]] | None = None,
) -> tuple[str, bytes]:
    """Объединяет логику: бинарный шаблон или текстовый legacy."""
    if docx_bytes:
        import json

        payload = json.loads(payload_json) if payload_json else {}
        if not isinstance(payload, dict):
            payload = {}
        docx_in = docx_bytes
        if conditional_blocks:
            docx_in, _warnings = apply_docx_conditional_blocks(docx_in, payload, conditional_blocks)
        flat = {str(k): "" if v is None else str(v) for k, v in payload.items()}
        return merge_docx_placeholders(docx_in, flat)
    result = generate_docx(docx_template_body, bindings_json, rules_json, payload_json)
    return result.file_name, result.bytes_
